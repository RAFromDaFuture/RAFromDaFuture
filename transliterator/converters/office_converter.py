"""
Office Document-to-Markdown Converter

Faithfully converts Word (.docx) and Excel (.xlsx) documents into
structured Markdown. Preserves headings, paragraphs, tables, and lists.
No interpretation — just transliteration.
"""

import os


class OfficeConverter:
    """Converts Office documents (.docx, .xlsx) to clean Markdown."""

    SUPPORTED_EXTENSIONS = {".docx", ".xlsx"}

    @staticmethod
    def can_handle(file_path: str) -> bool:
        _, ext = os.path.splitext(file_path.lower())
        return ext in OfficeConverter.SUPPORTED_EXTENSIONS

    @staticmethod
    def convert(file_path: str) -> str:
        """Convert a Word or Excel file to Markdown."""
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        _, ext = os.path.splitext(file_path.lower())

        if ext == ".docx":
            return _convert_docx(file_path)
        elif ext == ".xlsx":
            return _convert_xlsx(file_path)
        else:
            raise ValueError(f"Unsupported Office format: {ext}")


def _convert_docx(file_path: str) -> str:
    """Convert a Word document to Markdown."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = Document(file_path)
    filename = os.path.basename(file_path)
    lines = []

    header = (
        f"---\n"
        f"source_type: docx\n"
        f"source_file: {filename}\n"
        f"---\n\n"
    )

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue

            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            if not text:
                lines.append("")
                continue

            # Map Word heading styles to Markdown headings
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.replace("Heading", "").strip())
                    level = min(level, 6)
                except ValueError:
                    level = 2
                lines.append(f"{'#' * level} {text}")
                lines.append("")
            elif style_name == "Title":
                lines.append(f"# {text}")
                lines.append("")
            elif style_name.startswith("List"):
                # Detect numbered vs bullet lists
                if any(c.isdigit() for c in text[:3]):
                    lines.append(f"1. {text}")
                else:
                    lines.append(f"- {text}")
            elif para.runs and para.runs[0].bold and len(text) < 100:
                lines.append(f"**{text}**")
                lines.append("")
            else:
                lines.append(text)
                lines.append("")

        elif tag == "tbl":
            table = None
            for t in doc.tables:
                if t._element is element:
                    table = t
                    break
            if table:
                lines.append(_table_to_markdown(table))
                lines.append("")

    return header + "\n".join(lines).strip() + "\n"


def _convert_xlsx(file_path: str) -> str:
    """Convert an Excel spreadsheet to Markdown tables."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise RuntimeError("openpyxl is not installed. Run: pip install openpyxl")

    wb = load_workbook(file_path, read_only=True, data_only=True)
    filename = os.path.basename(file_path)

    header = (
        f"---\n"
        f"source_type: xlsx\n"
        f"source_file: {filename}\n"
        f"sheets: {', '.join(wb.sheetnames)}\n"
        f"---\n\n"
    )

    sections = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))

        if not rows:
            sections.append(f"## {sheet_name}\n\n[Empty sheet]\n")
            continue

        # Find the actual data bounds (skip fully empty rows)
        data_rows = []
        for row in rows:
            if any(cell is not None for cell in row):
                data_rows.append([str(cell) if cell is not None else "" for cell in row])

        if not data_rows:
            sections.append(f"## {sheet_name}\n\n[Empty sheet]\n")
            continue

        # Build markdown table
        section = f"## {sheet_name}\n\n"

        # Header row
        header_row = data_rows[0]
        section += "| " + " | ".join(header_row) + " |\n"
        section += "| " + " | ".join(["---"] * len(header_row)) + " |\n"

        # Data rows
        for row in data_rows[1:]:
            # Pad row if shorter than header
            padded = row + [""] * (len(header_row) - len(row))
            section += "| " + " | ".join(padded[:len(header_row)]) + " |\n"

        sections.append(section)

    wb.close()
    return header + "\n".join(sections).strip() + "\n"


def _table_to_markdown(table) -> str:
    """Convert a docx table to a Markdown table."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    # Determine column count from widest row
    col_count = max(len(r) for r in rows)
    for r in rows:
        while len(r) < col_count:
            r.append("")

    md = "| " + " | ".join(rows[0]) + " |\n"
    md += "| " + " | ".join(["---"] * col_count) + " |\n"
    for row in rows[1:]:
        md += "| " + " | ".join(row) + " |\n"

    return md
