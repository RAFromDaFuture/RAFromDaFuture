#!/usr/bin/env python3
"""
Transliterator - Faithful Document-to-Markdown Converter

A single-file app that converts PDFs, web pages, SharePoint links,
images (OCR), and Office docs (docx/xlsx) into clean structured Markdown.
Zero AI interpretation — just precise structural transliteration.

Usage:
    python transliterator_app.py policy.pdf
    python transliterator_app.py https://company.sharepoint.com/sites/HR/policy.aspx
    python transliterator_app.py ./documents/
    python transliterator_app.py file1.pdf file2.docx --stdout
    python transliterator_app.py --formats

Requirements:
    pip install pymupdf4llm pymupdf beautifulsoup4 markdownify requests pytesseract Pillow python-docx openpyxl
"""

import argparse
import os
import re
import sys
import tempfile
from datetime import datetime, timezone
from urllib.parse import urlparse

# ──────────────────────────────────────────────────────────────
# PDF CONVERTER
# ──────────────────────────────────────────────────────────────

class PDFConverter:
    SUPPORTED = {".pdf"}

    @staticmethod
    def can_handle(path):
        _, ext = os.path.splitext(path.lower())
        return ext in PDFConverter.SUPPORTED

    @staticmethod
    def convert(path):
        if not os.path.isfile(path):
            raise FileNotFoundError(f"PDF not found: {path}")

        try:
            import pymupdf4llm
            md = pymupdf4llm.to_markdown(path)
            return _meta("pdf", os.path.basename(path)) + _clean(md)
        except ImportError:
            pass

        try:
            import fitz
            doc = fitz.open(path)
            pages = []
            for i, page in enumerate(doc):
                text = page.get_text("text")
                if text.strip():
                    pages.append(f"<!-- Page {i+1} -->\n\n{text.strip()}")
            doc.close()
            return _meta("pdf", os.path.basename(path)) + _clean("\n\n---\n\n".join(pages))
        except ImportError:
            raise RuntimeError("Install pymupdf4llm or pymupdf:  pip install pymupdf4llm pymupdf")


# ──────────────────────────────────────────────────────────────
# WEB / SHAREPOINT CONVERTER
# ──────────────────────────────────────────────────────────────

class WebConverter:
    @staticmethod
    def can_handle(source):
        try:
            p = urlparse(source)
            return p.scheme in ("http", "https") and bool(p.netloc)
        except Exception:
            return False

    @staticmethod
    def is_sharepoint(url):
        host = urlparse(url).netloc.lower()
        return "sharepoint.com" in host or "sharepoint.us" in host

    @staticmethod
    def convert(url):
        import requests
        from bs4 import BeautifulSoup
        from markdownify import markdownify as md

        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/120.0.0.0 Safari/537.36"
            )
        }

        resp = requests.get(url, headers=headers, timeout=30, allow_redirects=True)
        resp.raise_for_status()
        ctype = resp.headers.get("Content-Type", "")

        # Remote PDF → hand off
        if "application/pdf" in ctype or url.lower().endswith(".pdf"):
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(resp.content)
                tmp_path = tmp.name
            try:
                return PDFConverter.convert(tmp_path)
            finally:
                os.unlink(tmp_path)

        # Remote Office doc → hand off
        if "application/vnd.openxmlformats-officedocument" in ctype:
            suffix = ".docx" if "wordprocessingml" in ctype else ".xlsx"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                tmp.write(resp.content)
                tmp_path = tmp.name
            try:
                return OfficeConverter.convert(tmp_path)
            finally:
                os.unlink(tmp_path)

        # Parse HTML
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup.find_all(["script", "style", "nav", "footer", "iframe", "noscript"]):
            tag.decompose()

        # SharePoint: zoom into main content
        if WebConverter.is_sharepoint(url):
            area = (
                soup.find("div", {"class": re.compile(r"CanvasZone|rte-editor|wiki", re.I)})
                or soup.find("div", {"role": "main"})
                or soup.find("main")
                or soup.find("article")
            )
            if area:
                soup = area

        # Title
        title = ""
        title_tag = soup.find("title") if hasattr(soup, "find") else None
        if title_tag and title_tag.string:
            title = title_tag.string.strip()
        elif soup.find("h1"):
            title = soup.find("h1").get_text(strip=True)

        md_text = md(
            str(soup),
            heading_style="ATX",
            bullets="-",
            strip=["img"],
            convert=["table", "thead", "tbody", "tr", "th", "td",
                     "p", "h1", "h2", "h3", "h4", "h5", "h6",
                     "ul", "ol", "li", "a", "strong", "em",
                     "blockquote", "pre", "code", "br", "hr"],
        )

        parsed = urlparse(url)
        stype = "sharepoint" if WebConverter.is_sharepoint(url) else "web"
        hdr = (
            f"---\n"
            f"source_type: {stype}\n"
            f"source_url: {url}\n"
            f"source_domain: {parsed.netloc}\n"
        )
        if title:
            hdr += f"title: {title}\n"
        hdr += "---\n\n"

        if title and not md_text.strip().startswith(f"# {title}"):
            md_text = f"# {title}\n\n{md_text}"

        return hdr + _clean(md_text)


# ──────────────────────────────────────────────────────────────
# IMAGE CONVERTER (OCR)
# ──────────────────────────────────────────────────────────────

class ImageConverter:
    SUPPORTED = {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".gif", ".webp"}

    @staticmethod
    def can_handle(path):
        _, ext = os.path.splitext(path.lower())
        return ext in ImageConverter.SUPPORTED

    @staticmethod
    def convert(path):
        if not os.path.isfile(path):
            raise FileNotFoundError(f"Image not found: {path}")

        from PIL import Image
        img = Image.open(path)
        w, h = img.size
        fmt = img.format or os.path.splitext(path)[1].lstrip(".")

        ocr_text = ""
        try:
            import pytesseract
            ocr_text = pytesseract.image_to_string(img)
        except ImportError:
            ocr_text = "[OCR unavailable — install pytesseract and Tesseract engine]"
        except Exception as e:
            ocr_text = f"[OCR failed: {e}]"

        hdr = (
            f"---\n"
            f"source_type: image\n"
            f"source_file: {os.path.basename(path)}\n"
            f"image_format: {fmt}\n"
            f"image_dimensions: {w}x{h}\n"
            f"---\n\n"
        )
        body = f"## Extracted Text\n\n{ocr_text.strip()}\n" if ocr_text.strip() else "## Extracted Text\n\n[No text detected in image]\n"
        return hdr + body


# ──────────────────────────────────────────────────────────────
# OFFICE CONVERTER (DOCX / XLSX)
# ──────────────────────────────────────────────────────────────

class OfficeConverter:
    SUPPORTED = {".docx", ".xlsx"}

    @staticmethod
    def can_handle(path):
        _, ext = os.path.splitext(path.lower())
        return ext in OfficeConverter.SUPPORTED

    @staticmethod
    def convert(path):
        if not os.path.isfile(path):
            raise FileNotFoundError(f"File not found: {path}")
        _, ext = os.path.splitext(path.lower())
        if ext == ".docx":
            return OfficeConverter._docx(path)
        elif ext == ".xlsx":
            return OfficeConverter._xlsx(path)
        raise ValueError(f"Unsupported format: {ext}")

    @staticmethod
    def _docx(path):
        from docx import Document
        doc = Document(path)
        lines = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                para = next((p for p in doc.paragraphs if p._element is element), None)
                if para is None:
                    continue
                style = para.style.name if para.style else ""
                text = para.text.strip()
                if not text:
                    lines.append("")
                    continue
                if style.startswith("Heading"):
                    try:
                        lvl = min(int(style.replace("Heading", "").strip()), 6)
                    except ValueError:
                        lvl = 2
                    lines.append(f"{'#' * lvl} {text}")
                    lines.append("")
                elif style == "Title":
                    lines.append(f"# {text}")
                    lines.append("")
                elif style.startswith("List"):
                    lines.append(f"- {text}")
                elif para.runs and para.runs[0].bold and len(text) < 100:
                    lines.append(f"**{text}**")
                    lines.append("")
                else:
                    lines.append(text)
                    lines.append("")

            elif tag == "tbl":
                table = next((t for t in doc.tables if t._element is element), None)
                if table:
                    lines.append(_table_to_md(table))
                    lines.append("")

        return _meta("docx", os.path.basename(path)) + "\n".join(lines).strip() + "\n"

    @staticmethod
    def _xlsx(path):
        from openpyxl import load_workbook
        wb = load_workbook(path, read_only=True, data_only=True)
        sections = []

        for name in wb.sheetnames:
            ws = wb[name]
            rows = [r for r in ws.iter_rows(values_only=True) if any(c is not None for c in r)]
            if not rows:
                sections.append(f"## {name}\n\n[Empty sheet]\n")
                continue

            data = [[str(c) if c is not None else "" for c in r] for r in rows]
            hdr = data[0]
            table = "| " + " | ".join(hdr) + " |\n"
            table += "| " + " | ".join(["---"] * len(hdr)) + " |\n"
            for row in data[1:]:
                padded = (row + [""] * len(hdr))[:len(hdr)]
                table += "| " + " | ".join(padded) + " |\n"
            sections.append(f"## {name}\n\n{table}")

        wb.close()
        hdr_meta = (
            f"---\n"
            f"source_type: xlsx\n"
            f"source_file: {os.path.basename(path)}\n"
            f"sheets: {', '.join(wb.sheetnames) if hasattr(wb, 'sheetnames') else ''}\n"
            f"---\n\n"
        )
        return hdr_meta + "\n".join(sections).strip() + "\n"


# ──────────────────────────────────────────────────────────────
# HELPERS
# ──────────────────────────────────────────────────────────────

def _meta(source_type, filename):
    return f"---\nsource_type: {source_type}\nsource_file: {filename}\n---\n\n"


def _clean(text):
    lines = text.split("\n")
    out, blanks = [], 0
    for line in lines:
        if line.strip() == "":
            blanks += 1
            if blanks <= 2:
                out.append(line)
        else:
            blanks = 0
            out.append(line)
    return "\n".join(out).strip() + "\n"


def _table_to_md(table):
    rows = [[cell.text.strip().replace("\n", " ") for cell in r.cells] for r in table.rows]
    if not rows:
        return ""
    cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < cols:
            r.append("")
    md = "| " + " | ".join(rows[0]) + " |\n"
    md += "| " + " | ".join(["---"] * cols) + " |\n"
    for row in rows[1:]:
        md += "| " + " | ".join(row) + " |\n"
    return md


def _safe_name(path):
    name, _ = os.path.splitext(os.path.basename(path))
    return "".join(c if c.isalnum() or c in "-_ " else "_" for c in name) + ".md"


def _url_name(url):
    p = urlparse(url)
    path = p.path.strip("/").replace("/", "_") or "index"
    domain = p.netloc.replace(".", "_")
    safe = "".join(c if c.isalnum() or c in "-_" else "_" for c in f"{domain}_{path}")
    return f"{safe}.md"


def _read_text(path):
    filename = os.path.basename(path)
    _, ext = os.path.splitext(filename)
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    hdr = _meta("text", filename)
    if ext.lower() in (".md", ".markdown"):
        return hdr + content
    lang_map = {
        ".py": "python", ".js": "javascript", ".ts": "typescript",
        ".json": "json", ".xml": "xml", ".html": "html", ".css": "css",
        ".yaml": "yaml", ".yml": "yaml", ".toml": "toml", ".sh": "bash",
        ".sql": "sql", ".csv": "csv",
    }
    lang = lang_map.get(ext.lower(), "")
    if lang:
        return hdr + f"```{lang}\n{content}\n```\n"
    return hdr + content


# ──────────────────────────────────────────────────────────────
# MAIN ENGINE
# ──────────────────────────────────────────────────────────────

ALL_FILE_EXTS = PDFConverter.SUPPORTED | ImageConverter.SUPPORTED | OfficeConverter.SUPPORTED


def convert_one(source, output_dir=None, to_stdout=False):
    """Convert a single source (file or URL) to markdown."""
    source = source.strip()

    if WebConverter.can_handle(source):
        print(f"  [URL] {source}")
        md = WebConverter.convert(source)
        name = _url_name(source)
    elif os.path.isfile(source):
        _, ext = os.path.splitext(source.lower())
        if PDFConverter.can_handle(source):
            print(f"  [PDF] {source}")
            md = PDFConverter.convert(source)
        elif ImageConverter.can_handle(source):
            print(f"  [IMG] {source}")
            md = ImageConverter.convert(source)
        elif OfficeConverter.can_handle(source):
            print(f"  [{ext.upper().lstrip('.')}] {source}")
            md = OfficeConverter.convert(source)
        else:
            print(f"  [TXT] {source}")
            md = _read_text(source)
        name = _safe_name(source)
    else:
        raise ValueError(f"Cannot handle: {source}")

    if to_stdout:
        print(md)
    elif output_dir:
        os.makedirs(output_dir, exist_ok=True)
        out = os.path.join(output_dir, name)
        with open(out, "w", encoding="utf-8") as f:
            f.write(md)
        print(f"  [SAVED] {out}")

    return md


def convert_dir(dir_path, output_dir=None, to_stdout=False):
    """Convert all supported files in a directory."""
    count = 0
    for fn in sorted(os.listdir(dir_path)):
        fp = os.path.join(dir_path, fn)
        if not os.path.isfile(fp):
            continue
        _, ext = os.path.splitext(fn.lower())
        if ext not in ALL_FILE_EXTS:
            continue
        try:
            convert_one(fp, output_dir=output_dir, to_stdout=to_stdout)
            count += 1
        except Exception as e:
            print(f"  [ERROR] {fn}: {e}", file=sys.stderr)
    return count


# ──────────────────────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        prog="transliterator",
        description=(
            "Faithful Document-to-Markdown Transliterator\n\n"
            "Converts PDFs, web pages, SharePoint links, images, and Office\n"
            "documents into clean Markdown. No AI interpretation — just\n"
            "precise structural transliteration."
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Examples:\n"
            "  python transliterator_app.py policy.pdf\n"
            "  python transliterator_app.py https://company.sharepoint.com/sites/HR/policy.aspx\n"
            "  python transliterator_app.py ./policies/\n"
            "  python transliterator_app.py doc1.pdf doc2.docx image.png\n"
            "  python transliterator_app.py report.pdf --stdout\n"
            "  python transliterator_app.py report.pdf -o ./md_output\n"
        ),
    )
    parser.add_argument("sources", nargs="*", help="Files, directories, or URLs to convert")
    parser.add_argument("-o", "--output", default="./transliterator_output", help="Output directory")
    parser.add_argument("--stdout", action="store_true", help="Print to stdout instead of saving")
    parser.add_argument("--formats", action="store_true", help="Show supported formats")

    args = parser.parse_args()

    if args.formats:
        print("\nSupported Input Formats:")
        print("-" * 40)
        print("  PDF:         .pdf")
        print("  Word:        .docx")
        print("  Excel:       .xlsx")
        print("  Images:      .png .jpg .jpeg .tiff .tif .bmp .gif .webp")
        print("  Web Pages:   http:// https://")
        print("  SharePoint:  *.sharepoint.com/*")
        print("  Plain Text:  .txt .csv .md .json .xml .html .py .js etc.")
        print()
        return

    if not args.sources:
        parser.print_help()
        print("\nError: No sources provided.")
        sys.exit(1)

    out_dir = None if args.stdout else args.output

    print("=" * 60)
    print("  TRANSLITERATOR — Faithful Document-to-Markdown Converter")
    print("=" * 60)
    print()

    ok, err = 0, 0
    for src in args.sources:
        try:
            if os.path.isdir(src):
                ok += convert_dir(src, output_dir=out_dir, to_stdout=args.stdout)
            else:
                convert_one(src, output_dir=out_dir, to_stdout=args.stdout)
                ok += 1
        except Exception as e:
            print(f"  [ERROR] {src}: {e}", file=sys.stderr)
            err += 1

    print()
    print("-" * 60)
    print(f"  Done: {ok} converted, {err} errors")
    if out_dir:
        print(f"  Output: {os.path.abspath(out_dir)}")
    print("-" * 60)


if __name__ == "__main__":
    main()
